"""
Generate readable DOCX from CorePAM LaTeX manuscript.
Converts LaTeX to plain text with figures embedded.
For personal review before submission.
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))

def resolve_labels(tex_content):
    """Build a map of label -> readable name from the LaTeX source."""
    labels = {}
    # Tables: find \caption{...}\label{...} patterns
    for m in re.finditer(r'\\caption\{[^}]*?\}\\label\{([^}]+)\}', tex_content):
        label = m.group(1)
        # Extract context from caption
        cap_block = m.group(0)
        if 'Table' in cap_block or 'tab:' in label:
            # Find table number from context
            if 'tab:survival' in label: labels[label] = 'Table 1'
            elif 'tab:coverage' in label: labels[label] = 'Table 2'
            elif 'tab:incremental' in label: labels[label] = 'Table 3'
            else: labels[label] = label.replace('tab:', 'Table ')
        elif 'fig:' in label:
            labels[label] = label.replace('fig:', 'Figure ')
    # Figures
    fig_labels = {
        'fig:study-design': 'Figure 1',
        'fig:km-multipanel': 'Figure 2',
        'fig:forest-meta': 'Figure 3',
        'fig:delta-c': 'Figure 4a',
        'fig:calibration': 'Figure 4b',
        'fig:pcr-forest': 'Figure 5',
    }
    labels.update(fig_labels)
    # Sections
    for m in re.finditer(r'\\(?:sub)?section\{([^}]+)\}\\label\{([^}]+)\}', tex_content):
        labels[m.group(2)] = m.group(1)
    return labels


def clean_latex(text, label_map=None):
    """Strip LaTeX commands to readable plain text."""
    if label_map is None:
        label_map = {}
    # Remove comments
    text = re.sub(r'%.*?$', '', text, flags=re.MULTILINE)
    # Remove document class, packages, etc.
    text = re.sub(r'\\documentclass.*?\n', '', text)
    text = re.sub(r'\\usepackage.*?\n', '', text)
    text = re.sub(r'\\unnumbered.*?\n', '', text)
    text = re.sub(r'\\raggedbottom', '', text)
    text = re.sub(r'\\begin\{document\}', '', text)
    text = re.sub(r'\\end\{document\}', '', text)
    text = re.sub(r'\\maketitle', '', text)
    # Handle special commands
    text = re.sub(r'\\title\[.*?\]\{(.*?)\}', r'TITLE: \1', text)
    text = re.sub(r'\\author\*?\[.*?\]\{.*?\\fnm\{(.*?)\}.*?\\sur\{(.*?)\}\}', r'Author: \1 \2', text)
    text = re.sub(r'\\email\{.*?\}', '', text)
    text = re.sub(r'\\affil\*?\[.*?\]\{.*?\}', '', text, flags=re.DOTALL)
    # Citations — replace with bracketed numbers
    text = re.sub(r'~?\\cite\{[^}]*\}', '', text)
    # References — resolve to readable names
    def ref_replacer(m):
        label = m.group(1)
        return label_map.get(label, label.split(':')[-1])
    text = re.sub(r'\\ref\{([^}]*)\}', ref_replacer, text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)
    # Formatting
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\fnm\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\sur\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\spfx\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\orgdiv\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\orgname\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\orgaddress\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\city\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\country\{([^}]*)\}', r'\1', text)
    # Math
    text = re.sub(r'\$\\Delta\$', 'Δ', text)
    text = re.sub(r'\$\\alpha\$', 'α', text)
    text = re.sub(r'\$\\beta\$', 'β', text)
    text = re.sub(r'\$\\tau\^2\$', 'τ²', text)
    text = re.sub(r'\$\\lambda\$', 'λ', text)
    text = re.sub(r'\$\\times\$', '×', text)
    text = re.sub(r'\$\\geq\$', '≥', text)
    text = re.sub(r'\$\\leq\$', '≤', text)
    text = re.sub(r'\$\\pm\$', '±', text)
    text = re.sub(r'\$\\approx\$', '≈', text)
    text = re.sub(r'\$I\^2\$', 'I²', text)
    text = re.sub(r'\$I\^{?2}?\$', 'I²', text)
    text = re.sub(r'\$p\$', 'p', text)
    text = re.sub(r'\$r\$', 'r', text)
    text = re.sub(r'\$B\s*=\s*(\d+)\$', r'B = \1', text)
    text = re.sub(r'\$K\s*=\s*(\d+)\$', r'K = \1', text)
    text = re.sub(r'\$k\s*=\s*(\d+)\$', r'k = \1', text)
    text = re.sub(r'\$N\s*=\s*', 'N = ', text)
    text = re.sub(r'\$<\$', '<', text)
    text = re.sub(r'\$>\$', '>', text)
    text = re.sub(r'\$\-\$', '−', text)
    text = re.sub(r'\$\+\$', '+', text)
    text = re.sub(r'\$([^$]*)\$', r'\1', text)  # strip remaining $...$
    # LaTeX special chars
    text = text.replace('\\&', '&')
    text = text.replace('\\%', '%')
    text = text.replace('\\#', '#')
    text = text.replace('\\\\', '\n')
    text = text.replace('~', ' ')
    text = text.replace('---', '—')
    text = text.replace('--', '–')
    text = text.replace("\\'{i}", 'í')
    text = text.replace("\\~{a}", 'ã')
    text = text.replace("\\'{e}", 'é')
    text = text.replace('\\_', '_')
    text = text.replace('\\,', ' ')
    # Sections
    text = re.sub(r'\\section\{([^}]*)\}', r'\n\n## \1\n', text)
    text = re.sub(r'\\subsection\{([^}]*)\}', r'\n### \1\n', text)
    # Lists
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\item', '•', text)
    # Abstract
    text = re.sub(r'\\abstract\{', '\n## Abstract\n', text)
    text = re.sub(r'\\keywords\{([^}]*)\}', r'\nKeywords: \1\n', text)
    # Misc
    text = re.sub(r'\\bmhead\{([^}]*)\}', r'\n## \1\n', text)
    text = re.sub(r'\\footnotetext(?:\[\d+\])?\{([^}]*)\}', r'[Note: \1]', text)
    text = re.sub(r'\\footnotemark\[\d+\]', '', text)
    # Remove remaining LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])*(?:\{[^}]*\})*', '', text)
    # Clean up
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def extract_figures(tex_content):
    """Extract figure paths and captions from LaTeX."""
    figs = []
    pattern = r'\\includegraphics(?:\[.*?\])?\{([^}]+)\}'
    caption_pattern = r'\\caption\{(.*?)\}(?:\\label\{.*?\})?'

    # Find all figure environments
    fig_envs = re.finditer(
        r'\\begin\{figure\}.*?\\end\{figure\}',
        tex_content, re.DOTALL
    )
    for m in fig_envs:
        block = m.group()
        img_m = re.search(pattern, block)
        cap_m = re.search(caption_pattern, block, re.DOTALL)
        if img_m:
            path = img_m.group(1)
            caption = clean_latex(cap_m.group(1)) if cap_m else ''
            figs.append({
                'path': path,
                'caption': caption,
                'pos': m.start()
            })
    return figs


def build_docx(tex_path, output_path):
    """Build a readable DOCX from LaTeX source."""
    with open(tex_path, 'r', encoding='utf-8') as f:
        tex = f.read()

    # Resolve cross-references
    label_map = resolve_labels(tex)

    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Title
    title_m = re.search(r'\\title\[.*?\]\{(.*?)\}', tex)
    if title_m:
        p = doc.add_heading(clean_latex(title_m.group(1), label_map), level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    doc.add_paragraph('Rafael de Negreiros Botan¹, João Batista de Sousa²')
    doc.add_paragraph('¹ Department of Oncology, Universidade de Brasília, Brazil')
    doc.add_paragraph('² Department of Proctology, Universidade de Brasília, Brazil')
    doc.add_paragraph('')

    # Extract figures for inline insertion
    figures = extract_figures(tex)
    fig_paths = {f['path']: f for f in figures}

    # Process sections
    # Split by major sections
    sections = re.split(r'(\\section\{[^}]*\})', tex)

    # Abstract first
    abs_m = re.search(r'\\abstract\{(.*?)\}\s*\\keywords', tex, re.DOTALL)
    if abs_m:
        doc.add_heading('Abstract', level=1)
        abs_text = clean_latex(abs_m.group(1), label_map)
        for para in abs_text.split('\n\n'):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    kw_m = re.search(r'\\keywords\{(.*?)\}', tex)
    if kw_m:
        doc.add_paragraph(f'Keywords: {clean_latex(kw_m.group(1), label_map)}')

    # Main content - extract between \maketitle and \end{document}
    main_m = re.search(r'\\maketitle(.*?)\\end\{document\}', tex, re.DOTALL)
    if not main_m:
        main_m = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', tex, re.DOTALL)

    if main_m:
        content = main_m.group(1)

        # Remove abstract (already handled)
        content = re.sub(r'\\abstract\{.*?\}\s*\\keywords\{.*?\}', '', content, flags=re.DOTALL)

        # Process figure environments - replace with markers
        fig_counter = [0]
        def fig_replacer(m):
            block = m.group()
            img_m = re.search(r'\\includegraphics(?:\[.*?\])?\{([^}]+)\}', block)
            cap_m = re.search(r'\\caption\{(.*?)\}', block, re.DOTALL)
            if img_m:
                fig_counter[0] += 1
                path = img_m.group(1)
                caption = clean_latex(cap_m.group(1), label_map) if cap_m else ''
                return f'\n\n[FIGURE:{path}:{caption}]\n\n'
            return ''

        content = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', fig_replacer, content, flags=re.DOTALL)

        # Process table environments - extract content
        def table_replacer(m):
            block = m.group()
            cap_m = re.search(r'\\caption\{(.*?)\}', block, re.DOTALL)
            caption = clean_latex(cap_m.group(1), label_map) if cap_m else 'Table'

            # Extract tabular content
            tab_m = re.search(r'\\begin\{tabular[*]?\}.*?\n(.*?)\\end\{tabular[*]?\}', block, re.DOTALL)
            if tab_m:
                rows = tab_m.group(1)
                rows = re.sub(r'\\toprule|\\midrule|\\botrule|\\bottomrule|\\hline', '', rows)
                rows = re.sub(r'\\\\', '\n', rows)
                rows = [clean_latex(r.strip(), label_map) for r in rows.split('\n') if r.strip() and not r.strip().startswith('\\')]
                table_text = '\n'.join(rows)
                return f'\n\n[TABLE: {caption}]\n{table_text}\n\n'
            return f'\n\n[TABLE: {caption}]\n\n'

        content = re.sub(r'\\begin\{table\}.*?\\end\{table\}', table_replacer, content, flags=re.DOTALL)

        # Clean the text
        cleaned = clean_latex(content, label_map)

        # Split into paragraphs and process
        paragraphs = cleaned.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Section headers
            if para.startswith('## '):
                doc.add_heading(para[3:].strip(), level=1)
            elif para.startswith('### '):
                doc.add_heading(para[4:].strip(), level=2)

            # Figure markers
            elif para.startswith('[FIGURE:'):
                m = re.match(r'\[FIGURE:(.*?):(.*?)\]', para, re.DOTALL)
                if m:
                    fig_path = m.group(1)
                    caption = m.group(2)

                    # Try to find the actual file
                    full_path = os.path.join(BASE, fig_path)
                    if not os.path.exists(full_path):
                        # Try without extension change
                        for ext in ['.png', '.pdf', '.jpg']:
                            alt = os.path.splitext(full_path)[0] + ext
                            if os.path.exists(alt):
                                full_path = alt
                                break

                    if os.path.exists(full_path) and full_path.lower().endswith('.png'):
                        try:
                            doc.add_picture(full_path, width=Inches(6))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            doc.add_paragraph(f'[Figure: {fig_path} — could not embed: {e}]')
                    else:
                        doc.add_paragraph(f'[Figure: {fig_path}]')

                    if caption:
                        p = doc.add_paragraph(caption)
                        p.style = doc.styles['Normal']
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.italic = True

            # Table markers
            elif para.startswith('[TABLE:'):
                doc.add_paragraph(para)

            # Regular text
            else:
                doc.add_paragraph(para)

    doc.save(output_path)
    print(f'DOCX saved: {output_path}')


if __name__ == '__main__':
    tex_path = os.path.join(BASE, 'CorePAM_manuscript.tex')
    out_path = os.path.join(BASE, 'CorePAM_manuscript_v3_review.docx')
    build_docx(tex_path, out_path)
