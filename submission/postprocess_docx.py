#!/usr/bin/env python3
"""
Post-process CorePAM_manuscript.docx to match Springer BCR formatting
(replicating the old submitted article structure).
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT  = "CorePAM_manuscript.docx"
OUTPUT = "CorePAM_manuscript.docx"

doc = Document(INPUT)

# =========================================================================
# Helper functions
# =========================================================================
def set_double_spacing(para):
    """Force double line spacing (480 twips) on a paragraph."""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="480" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn('w:line'), '480')
        spacing.set(qn('w:lineRule'), 'auto')

def set_style_by_xml(para, style_id):
    """Set paragraph style via raw XML (more reliable than python-docx style assignment)."""
    pPr = para._element.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = parse_xml(f'<w:pStyle {nsdecls("w")} w:val="{style_id}"/>')
        pPr.insert(0, pStyle)
    else:
        pStyle.set(qn('w:val'), style_id)

def add_run(para, text, bold=False, font_name='Times New Roman', font_size=None):
    """Add a run with formatting."""
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    return run

def insert_para_after(ref_para):
    """Insert an empty paragraph element after ref_para and return it as Paragraph."""
    new_el = parse_xml(f'<w:p {nsdecls("w")}/>')
    ref_para._element.addnext(new_el)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_el, ref_para._element.getparent())

# =========================================================================
# 0. Remove empty Author paragraphs left by pandoc
# =========================================================================
to_remove = []
for p in doc.paragraphs:
    if p.style.name == 'Author' and p.text.strip() == '':
        to_remove.append(p)
for p in to_remove:
    p._element.getparent().remove(p._element)

# =========================================================================
# 1. Insert title page content after paragraph[0] (Title)
# =========================================================================
title_para = doc.paragraphs[0]

# Build chain: title → authors → affiliations → corresponding → [page break] → abstract → ...
# Insert in forward order using insert_para_after

# --- Authors ---
p = insert_para_after(title_para)
set_style_by_xml(p, 'FirstParagraph')
add_run(p, 'Rafael de Negreiros Botan\u00B9*, Jo\u00E3o Batista de Sousa\u00B2')
set_double_spacing(p)
p_authors = p

# --- Affiliations ---
p = insert_para_after(p_authors)
set_style_by_xml(p, 'FirstParagraph')
add_run(p, '\u00B9 Department of Oncology, Universidade de Bras\u00EDlia \u2013 Bras\u00EDlia, Brazil')
set_double_spacing(p)
p_affil1 = p

p = insert_para_after(p_affil1)
set_style_by_xml(p, 'FirstParagraph')
add_run(p, '\u00B2 Department of Proctology, Universidade de Bras\u00EDlia \u2013 Bras\u00EDlia, Brazil')
set_double_spacing(p)
p_affil2 = p

# --- Corresponding author ---
p = insert_para_after(p_affil2)
set_style_by_xml(p, 'FirstParagraph')
add_run(p, '*Corresponding author: Rafael de Negreiros Botan \u2013 oncologista@gmail.com')
set_double_spacing(p)
p_corresp = p

# --- Page break ---
p = insert_para_after(p_corresp)
run_el = parse_xml(f'<w:r {nsdecls("w")}><w:br w:type="page"/></w:r>')
p._element.append(run_el)
p_break = p

# --- Abstract heading ---
p = insert_para_after(p_break)
set_style_by_xml(p, 'Ttulo1')  # Heading 1 XML ID (Portuguese locale)
add_run(p, 'Abstract', font_name='Times New Roman')
set_double_spacing(p)
p_abstract_h = p

# --- Abstract structured paragraphs ---
abstract_sections = [
    ("Background: ",
     "The PAM50 classifier predicts breast cancer prognosis but requires 50 genes and specialised platforms. "
     "We derived CorePAM, the smallest data-driven PAM50 subset maintaining non-inferior prognostic performance "
     "relative to a full 50-gene Cox elastic-net model, without pre-specifying gene count."),
    ("Methods: ",
     "Cox elastic-net regression (\u03B1 = 0.5) with deterministic 10-fold cross-validation was applied in the "
     "SCAN-B cohort (N = 3,069; GSE96058). Gene selection followed a pre-specified non-inferiority margin "
     "(\u0394C-index = 0.010). External validation used four independent cohorts: TCGA-BRCA (N = 1,072; RNA-seq), "
     "METABRIC (N = 1,978; microarray; disease-specific survival), GSE20685 (N = 327; microarray), and GSE1456 "
     "(N = 159; microarray). Incremental value over a clinical model (CORE-A: age and ER status) was assessed by "
     "bootstrap \u0394C-index. Secondary analyses evaluated pathologic complete response (pCR) prediction in four "
     "neoadjuvant cohorts (N = 697) plus I-SPY2 (N = 986)."),
    ("Results: ",
     "CorePAM comprises 24 genes with an out-of-fold C-index of 0.670 (gap vs 50-gene maximum: 0.009). "
     "The score was independently associated with survival in all validation cohorts: TCGA-BRCA (HR = 1.20), "
     "METABRIC DSS (HR = 1.41), GSE20685 (HR = 1.40), GSE1456 (HR = 1.71); all p < 0.02. Random-effects "
     "meta-analysis (K = 4) yielded pooled HR = 1.37 (95% CI 1.24\u20131.52; I\u00B2 = 38.3%; "
     "p = 1.8 \u00D7 10\u207B\u2079). CorePAM provided incremental value beyond CORE-A and remained significant "
     "after adjustment for T-stage and nodal status. For pCR, pooled OR = 1.69 (95% CI 1.39\u20132.05; "
     "I\u00B2 = 0%; p = 1.9 \u00D7 10\u207B\u2077)."),
    ("Conclusions: ",
     "CorePAM\u2014a 24-gene PAM50-derived score\u2014achieves non-inferior discrimination relative to a 50-gene "
     "Cox elastic-net model across four validation cohorts spanning RNA-seq and microarray platforms, remains "
     "significant after anatomical staging adjustment, and predicts pCR. This reduction from 50 to 24 genes may "
     "facilitate implementation in resource-limited settings."),
]

prev = p_abstract_h
for bold_label, text in abstract_sections:
    p = insert_para_after(prev)
    set_style_by_xml(p, 'Normal')
    add_run(p, bold_label, bold=True)
    add_run(p, text, bold=False)
    set_double_spacing(p)
    prev = p

# --- Trial registration ---
p = insert_para_after(prev)
set_style_by_xml(p, 'Subttulo')  # Subtitle XML ID (Portuguese locale)
add_run(p, 'Trial registration: Not applicable.')
set_double_spacing(p)
prev = p

# --- Keywords ---
p = insert_para_after(prev)
set_style_by_xml(p, 'Subttulo')
add_run(p, 'Keywords')
set_double_spacing(p)
prev = p

p = insert_para_after(prev)
set_style_by_xml(p, 'FirstParagraph')
add_run(p, 'Breast cancer; Gene expression; Prognostic signature; PAM50; Elastic-net regression; '
           'Survival analysis; Pathologic complete response; Cross-platform validation')
set_double_spacing(p)

# =========================================================================
# 2. Convert all Heading 2 → Subtitle
# =========================================================================
for p in doc.paragraphs:
    if p.style.name == 'Heading 2':
        set_style_by_xml(p, 'Subttulo')

# =========================================================================
# 3. Double spacing on ALL paragraphs
# =========================================================================
for p in doc.paragraphs:
    set_double_spacing(p)

# =========================================================================
# 4. Set Times New Roman on all runs without explicit font
# =========================================================================
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.name is None:
            run.font.name = 'Times New Roman'

# =========================================================================
# 5. Force double spacing on style definitions too
# =========================================================================
for style_name in ['Normal', 'Body Text', 'First Paragraph', 'Heading 1',
                    'Heading 2', 'Subtitle', 'Author', 'Title']:
    try:
        doc.styles[style_name].paragraph_format.line_spacing = 2.0
    except KeyError:
        pass

# =========================================================================
# Save
# =========================================================================
doc.save(OUTPUT)

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
print(f"OK: {OUTPUT}")
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Verify structure
for i, p in enumerate(doc.paragraphs[:20]):
    txt = p.text[:80] if p.text else '(empty)'
    print(f'  [{i:2d}] {p.style.name}: {txt}')
