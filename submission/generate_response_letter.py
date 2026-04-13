"""
Generate point-by-point response letter as DOCX for CorePAM R2 revision.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(os.path.abspath(__file__))


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_reviewer_comment(doc, text):
    """Add reviewer comment in italic gray."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(10)
    return p


def add_response(doc, text):
    """Add author response in normal style."""
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p


def add_bold_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def build_response_letter(output_path):
    doc = Document()

    # Style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Title
    p = doc.add_heading('Point-by-Point Response to Reviewers', level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Manuscript: CorePAM: a 24-gene PAM50-derived expression score with cross-platform external validation for breast cancer prognosis')
    doc.add_paragraph('Journal: Breast Cancer Research')
    doc.add_paragraph('Round: R2 Revision')
    doc.add_paragraph('')

    # ================================================================
    # REVIEWER 1
    # ================================================================
    add_heading_styled(doc, 'Reviewer 1', level=1)

    doc.add_paragraph('We thank the reviewer for their constructive feedback, which has improved the presentation of the manuscript.')
    doc.add_paragraph('')

    # R1.1
    add_bold_line(doc, 'Comment R1.1 — Kaplan-Meier figures')
    add_reviewer_comment(doc,
        'Figures 2-6 present Kaplan Meier curves for the different datasets. '
        'The annotations on the plots report two different p-values. Each plot should report '
        'a single p-value clearly linked to the comparison being presented (or, if multiple tests '
        'are intentionally reported, the figure legend should explicitly define what each p-value '
        'corresponds to). In addition, these figures could be consolidated into a single multi-panel '
        'figure (one dataset per panel). For consistency, all panels should use the same color scheme, '
        'legend format, and annotation style.')

    add_bold_line(doc, 'Response:')
    add_response(doc,
        'We thank the reviewer for this constructive suggestion. Figures 2\u20136 have been consolidated '
        'into a single multi-panel Figure 2 (panels A\u2013E, one per cohort). Each panel now '
        'reports a single log-rank p-value and the dichotomised Cox HR with 95% CI. All panels share a '
        'consistent colour scheme, legend format, and annotation style.')

    doc.add_paragraph('')

    # R1.2
    add_bold_line(doc, 'Comment R1.2 — Supplementary numbering')
    add_reviewer_comment(doc,
        'Supplementary figures and tables should be numbered in the order in which they are first '
        'cited in the text. For example, Figure S2 is currently cited before Figure S1, which is confusing.')

    add_bold_line(doc, 'Response:')
    add_response(doc,
        'All supplementary materials have been renumbered as Additional files 1\u201318 following the order '
        'of first citation in the text. Cross-references throughout the manuscript were updated accordingly.')

    doc.add_paragraph('')

    # R1.3
    add_bold_line(doc, 'Comment R1.3 — Unreferenced supplementary figures')
    add_reviewer_comment(doc,
        'All supplementary figures should be referenced in the main text. At present, '
        'Figures S3 and S4 are not mentioned.')

    add_bold_line(doc, 'Response:')
    add_response(doc,
        'All supplementary figures and tables are now explicitly referenced in the main text at the '
        'point of first relevance.')

    doc.add_paragraph('')

    # R1.4
    add_bold_line(doc, 'Comment R1.4 — Discussion narrative flow and limitations')
    add_reviewer_comment(doc,
        'The Discussion lacks narrative flow and largely mirrors the Results section, which makes it '
        'feel repetitive. The limitations section, in particular, reads as a list; it would benefit '
        'from explaining how the limitations may have affected the analyses/interpretation and, where '
        'possible, how they could be addressed in future work.')

    add_bold_line(doc, 'Response:')
    add_response(doc,
        'We appreciate this feedback and have revised the Discussion to improve narrative flow. '
        'The opening paragraph was rewritten to frame the principal finding and its scope rather than '
        'restating numerical results. A new paragraph presents the direct external head-to-head '
        'comparison between CorePAM and the fuller PAM50-based comparator. The Limitations section was '
        'restructured into three interpretive subsections\u2014study design and data, analytical choices, '
        'and clinical translation\u2014with commentary on how each limitation may affect interpretation.')

    doc.add_paragraph('')

    # ================================================================
    # REVIEWER 2
    # ================================================================
    add_heading_styled(doc, 'Reviewer 2', level=1)

    # R2
    add_bold_line(doc, 'Comment R2 — Tone down claims on cross-platform applicability, clinical utility, and non-inferiority')
    add_reviewer_comment(doc,
        'This revised manuscript is much improved, but a few issues still need revision: the claims '
        'on cross-platform applicability, clinical utility, and \u201cnon-inferiority\u201d should be toned down, '
        'because external calibration remains suboptimal and the frozen-reference analysis suggests that '
        'single-sample performance is not yet stable across all platforms.')

    add_bold_line(doc, 'Response:')
    add_response(doc,
        'We appreciate this important observation and revised the manuscript accordingly. We now '
        'distinguish external validation across RNA-seq and microarray cohorts from platform-agnostic '
        'single-sample deployment, and we moderated statements implying immediate clinical utility. '
        'We also clarified that the non-inferiority claim is anchored to the elastic-net Cox framework.')

    add_response(doc,
        'To further assess whether the 50-to-24 gene reduction compromised external discrimination, '
        'we added a direct paired comparison between CorePAM and the PAM50-full elastic-net model '
        'trained on all 50 candidate genes (49 non-zero coefficients at the selected \u03bb). In these paired '
        'bootstrap analyses, discrimination was similar in TCGA-BRCA and favoured CorePAM in METABRIC, '
        'GSE20685, and GSE1456 under both cohort-standardised and frozen-reference scoring '
        '(Additional file 18).')

    add_response(doc,
        'We agree that external calibration remained imperfect and that frozen-reference performance was '
        'not uniform across platforms; we therefore now frame these findings explicitly as a limitation '
        'to transportability, requiring platform-specific recalibration before deployment.')

    doc.add_paragraph('')

    # ================================================================
    # UNSOLICITED CORRECTIONS
    # ================================================================
    add_heading_styled(doc, 'Note on unsolicited corrections', level=1)

    add_response(doc,
        'During final quality-control review, we identified and corrected an incorrect platform '
        'assignment for GSE20685 (HG-U133A instead of HG-U133 Plus 2.0 / GPL570). After re-running '
        'the affected analyses, CorePAM coverage changed from 22/24 to 24/24, while the primary '
        'quantitative conclusions remained materially unchanged (largest change: +0.010 in frozen '
        'C-index; meta-analytic HR unchanged). All affected outputs included in the resubmission were '
        'regenerated. We also added GSE1456 to the frozen-reference sensitivity analysis for '
        'completeness.')

    doc.save(output_path)
    print(f'Response letter saved: {output_path}')


if __name__ == '__main__':
    out = os.path.join(BASE, 'CorePAM_PointByPoint_R2.docx')
    build_response_letter(out)
